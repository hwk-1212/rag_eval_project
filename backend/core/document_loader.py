import fitz  # PyMuPDF
from pathlib import Path
from typing import List, Dict, Any, Optional
from loguru import logger
import hashlib


class Document:
    """文档数据类"""
    def __init__(
        self,
        content: str,
        metadata: Optional[Dict[str, Any]] = None,
        doc_id: Optional[str] = None
    ):
        self.content = content
        self.metadata = metadata or {}
        self.doc_id = doc_id or self._generate_id()
    
    def _generate_id(self) -> str:
        """生成文档唯一ID"""
        return hashlib.md5(self.content.encode()).hexdigest()[:16]


class Chunk(Document):
    """文档块数据类"""
    def __init__(
        self,
        content: str,
        chunk_index: int,
        metadata: Optional[Dict[str, Any]] = None,
        chunk_id: Optional[str] = None
    ):
        super().__init__(content, metadata, chunk_id)
        self.chunk_index = chunk_index
        self.metadata["chunk_index"] = chunk_index


class DocumentLoader:
    """文档加载器 - 支持多种文件格式"""
    
    def __init__(self):
        self.supported_formats = {".pdf", ".txt", ".md", ".docx"}
    
    def load(self, file_path: str) -> Document:
        """
        加载文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            Document对象
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        file_ext = file_path.suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {file_ext}")
        
        try:
            if file_ext == ".pdf":
                text = self._load_pdf(file_path)
            elif file_ext in {".txt", ".md"}:
                text = self._load_text(file_path)
            elif file_ext == ".docx":
                text = self._load_docx(file_path)
            else:
                raise ValueError(f"未实现的文件格式处理: {file_ext}")
            
            metadata = {
                "filename": file_path.name,
                "file_type": file_ext,
                "file_path": str(file_path),
                "file_size": file_path.stat().st_size,
            }
            
            logger.info(f"成功加载文档: {file_path.name}, 长度: {len(text)} 字符")
            return Document(content=text, metadata=metadata)
            
        except Exception as e:
            logger.error(f"加载文档失败 {file_path}: {e}")
            raise
    
    def _load_pdf(self, file_path: Path) -> str:
        """从PDF提取文本"""
        doc = fitz.open(file_path)
        text = ""
        for page_num in range(doc.page_count):
            page = doc[page_num]
            text += page.get_text("text")
        doc.close()
        return text
    
    def _load_text(self, file_path: Path) -> str:
        """加载纯文本文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    
    def _load_docx(self, file_path: Path) -> str:
        """加载Word文档"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            raise ImportError("需要安装python-docx: pip install python-docx")
    
    def chunk(
        self,
        document: Document,
        chunk_size: int = 500,
        chunk_overlap: int = 100,
        strategy: str = "fixed"
    ) -> List[Chunk]:
        """
        文档分块
        
        Args:
            document: Document对象
            chunk_size: 块大小
            chunk_overlap: 重叠大小
            strategy: 分块策略 (fixed/semantic)
            
        Returns:
            Chunk对象列表
        """
        if strategy == "fixed":
            return self._fixed_chunk(document, chunk_size, chunk_overlap)
        elif strategy == "semantic":
            # TODO: 实现语义分块
            logger.warning("语义分块尚未实现，回退到固定长度分块")
            return self._fixed_chunk(document, chunk_size, chunk_overlap)
        else:
            raise ValueError(f"不支持的分块策略: {strategy}")
    
    def _fixed_chunk(
        self,
        document: Document,
        chunk_size: int,
        chunk_overlap: int
    ) -> List[Chunk]:
        """固定长度分块"""
        text = document.content
        chunks = []
        
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end]
            
            if not chunk_text.strip():
                start = end
                continue
            
            metadata = document.metadata.copy()
            metadata.update({
                "start_index": start,
                "end_index": end,
                "chunk_size": len(chunk_text)
            })
            
            chunk = Chunk(
                content=chunk_text,
                chunk_index=chunk_index,
                metadata=metadata
            )
            chunks.append(chunk)
            
            start += (chunk_size - chunk_overlap)
            chunk_index += 1
        
        logger.info(f"文档分块完成: {len(chunks)} 个块")
        return chunks
    
    def batch_load(self, file_paths: List[str]) -> List[Document]:
        """批量加载文档"""
        documents = []
        for file_path in file_paths:
            try:
                doc = self.load(file_path)
                documents.append(doc)
            except Exception as e:
                logger.error(f"加载文档失败 {file_path}: {e}")
                continue
        
        logger.info(f"批量加载完成: {len(documents)}/{len(file_paths)} 个文档")
        return documents

